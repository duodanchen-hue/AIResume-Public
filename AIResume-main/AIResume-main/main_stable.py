import html
import json
import os
import shutil
import uuid
import random
import string
from typing import Optional

import pdfplumber
import pytesseract
from PIL import Image
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import HTMLResponse, FileResponse, RedirectResponse
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi.responses import FileResponse

app = FastAPI()
from fastapi.staticfiles import StaticFiles

if os.path.exists("static"):
    app.mount("/static", StaticFiles(directory="static"), name="static")
PRO_CODES_FILE = "pro_codes.json"
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin123")


def load_codes():
    if not os.path.exists(PRO_CODES_FILE):
        return {}

    try:
        with open(PRO_CODES_FILE, "r", encoding="utf-8") as f:
            content = f.read().strip()
            if not content:
                return {}
            return json.loads(content)
    except Exception:
        return {}


def save_codes(codes):
    with open(PRO_CODES_FILE, "w", encoding="utf-8") as f:
        json.dump(codes, f, ensure_ascii=False, indent=2)


def generate_pro_code(length=10):
    chars = string.ascii_uppercase + string.digits
    return "AI-" + "".join(random.choices(chars, k=length))


def create_codes(count=10):
    codes = load_codes()
    new_codes = []

    for _ in range(count):
        code = generate_pro_code()
        while code in codes:
            code = generate_pro_code()

        codes[code] = {
            "used": False
        }
        new_codes.append(code)

    save_codes(codes)
    return new_codes


def is_pro_user(pro_code: str) -> bool:
    code = pro_code.strip()
    codes = load_codes()

    if code not in codes:
        return False

    if codes[code].get("used") is True:
        return False

    codes[code]["used"] = True
    save_codes(codes)

    return True

# Tesseract OCR 路径
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# 从环境变量读取 API Key，不要把真实 key 写进代码或上传到 GitHub。
# Windows PowerShell: $env:DEEPSEEK_API_KEY="你的key"
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")


def get_deepseek_client():
    if not DEEPSEEK_API_KEY:
        raise RuntimeError("缺少 DEEPSEEK_API_KEY 环境变量，请在本地或部署平台配置后再启动服务。")

    return OpenAI(
        api_key=DEEPSEEK_API_KEY,
        base_url="https://api.deepseek.com"
    )


def save_upload_file(upload_file: UploadFile) -> str:
    os.makedirs("uploads", exist_ok=True)
    ext = os.path.splitext(upload_file.filename or "")[1]
    filename = os.path.join("uploads", f"{uuid.uuid4().hex}{ext}")
    with open(filename, "wb") as buffer:
        shutil.copyfileobj(upload_file.file, buffer)
    return filename


def extract_pdf_text(file: UploadFile) -> str:
    file_path = save_upload_file(file)
    resume_text = ""

    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                resume_text += text + "\n"

    return resume_text.strip()


def extract_image_text(jd_image: Optional[UploadFile]) -> str:
    if not jd_image or not jd_image.filename:
        return ""

    image_path = save_upload_file(jd_image)
    image = Image.open(image_path)

    return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()


def strip_json_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```json"):
        text = text.replace("```json", "", 1).replace("```", "").strip()
    elif text.startswith("```"):
        text = text.replace("```", "").strip()
    return text


def e(value):
    return html.escape(str(value or ""))


def render_tags(items, bg, color):
    if not items:
        return "<p style='color:#9ca3af;'>暂无</p>"

    return "".join(
        f"<span style='display:inline-block;background:{bg};color:{color};padding:7px 12px;border-radius:999px;margin:4px;font-size:14px;font-weight:600;'>{e(item)}</span>"
        for item in items
    )


def render_list(items):
    if not items:
        return "<p style='color:#9ca3af;'>暂无</p>"

    return "<ul>" + "".join(
        f"<li>{e(item)}</li>" for item in items
    ) + "</ul>"


def list_items(items):
    if not items:
        return ""
    return "".join(f"<li>{e(item)}</li>" for item in items if item)


def skills_tags(items):
    if not items:
        return "<span class='muted'>暂无</span>"

    # 兼容两种格式：["Python"] 或 [{"category":"工具","items":["Python"]}]
    tags = ""
    for item in items:
        if isinstance(item, dict):
            category = item.get("category", "")
            sub_items = item.get("items", [])
            if category:
                tags += f"<div class='skill-category'>{e(category)}</div>"
            tags += "".join(f"<span class='skill-tag'>{e(x)}</span>" for x in sub_items if x)
        else:
            tags += f"<span class='skill-tag'>{e(item)}</span>"
    return tags


def plain_text_resume(data):
    lines = []
    lines.append(data.get("name", ""))
    contact = "｜".join([x for x in [
        data.get("phone", ""),
        data.get("email", ""),
        data.get("city", ""),
        data.get("target_job", "")
    ] if x])
    lines.append(contact)
    lines.append("")

    lines.append("个人简介")
    lines.append(data.get("summary", ""))
    lines.append("")

    lines.append("教育经历")
    for edu in data.get("education", []):
        lines.append(f"{edu.get('school','')}｜{edu.get('major','')}｜{edu.get('degree','')}｜{edu.get('time','')}")
        for d in edu.get("details", []):
            lines.append(f"- {d}")
    lines.append("")

    lines.append("技能优势")
    for s in data.get("skills", []):
        if isinstance(s, dict):
            category = s.get("category", "")
            items = "、".join(s.get("items", []))
            lines.append(f"- {category}：{items}" if category else f"- {items}")
        else:
            lines.append(f"- {s}")
    lines.append("")

    lines.append("实习经历 / 工作经历")
    for exp in data.get("experiences", []):
        lines.append(f"{exp.get('company','')}｜{exp.get('role','')}｜{exp.get('time','')}")
        for b in exp.get("bullets", []):
            lines.append(f"- {b}")
    lines.append("")

    lines.append("项目经历")
    for p in data.get("projects", []):
        title = p.get("name", "")
        role = p.get("role", "")
        time = p.get("time", "")
        lines.append("｜".join([x for x in [title, role, time] if x]))
        for b in p.get("bullets", []):
            lines.append(f"- {b}")
    lines.append("")

    suggested = data.get("suggested_projects", [])
    if suggested:
        lines.append("AI 推荐补做项目（完成后再写入简历）")
        for p in suggested:
            lines.append(f"{p.get('name','')}｜适合岗位：{p.get('suitable_for','')}｜难度：{p.get('difficulty','')}｜预计：{p.get('estimated_days','')}")
            if p.get("reason"):
                lines.append(f"建议原因：{p.get('reason')}")
            lines.append("需要真实完成：")
            for task in p.get("must_complete_tasks", []):
                lines.append(f"- {task}")
            lines.append("完成后可写入简历：")
            for b in p.get("resume_bullets_after_completion", []):
                lines.append(f"- {b}")
        lines.append("")

    lines.append("荣誉奖项")
    for a in data.get("awards", []):
        lines.append(f"- {a}")
    lines.append("")

    lines.append("其他经历")
    for o in data.get("others", []):
        lines.append(f"- {o}")

    return "\n".join(lines)


def build_resume_html(data):
    education_html = ""
    for edu in data.get("education", []):
        education_html += f"""
        <div class="item">
            <div class="item-title">{e(edu.get("school"))}</div>
            <div class="item-meta">{e(edu.get("major"))}｜{e(edu.get("degree"))}｜{e(edu.get("time"))}</div>
            <ul>{list_items(edu.get("details", []))}</ul>
        </div>
        """

    experience_html = ""
    for exp in data.get("experiences", []):
        experience_html += f"""
        <div class="item">
            <div class="item-title">{e(exp.get("company"))}｜{e(exp.get("role"))}</div>
            <div class="item-meta">{e(exp.get("time"))}</div>
            <ul>{list_items(exp.get("bullets", []))}</ul>
        </div>
        """

    projects_html = ""
    for p in data.get("projects", []):
        title = "｜".join([x for x in [p.get("name", ""), p.get("role", "")] if x])
        projects_html += f"""
        <div class="item">
            <div class="item-title">{e(title)}</div>
            <div class="item-meta">{e(p.get("time"))}</div>
            <ul>{list_items(p.get("bullets", []))}</ul>
        </div>
        """

    suggested_html = ""
    suggested = data.get("suggested_projects", [])
    if suggested:
        suggested_html += "<div class='section-title warning-title'>AI 推荐补做项目（完成后再写入简历）</div>"
        for p in suggested:
            suggested_html += f"""
            <div class="item suggested">
                <div class="item-title">{e(p.get("name"))}</div>
                <div class="item-meta">
                    适合岗位：{e(p.get("suitable_for"))}｜难度：{e(p.get("difficulty"))}｜预计：{e(p.get("estimated_days"))}
                </div>
                <div class="reason"><strong>推荐原因：</strong>{e(p.get("reason"))}</div>

                <div class="project-subtitle">你需要真实完成：</div>
                <ul>{list_items(p.get("must_complete_tasks", []))}</ul>

                <div class="project-subtitle">完成后可写入简历：</div>
                <ul>{list_items(p.get("resume_bullets_after_completion", []))}</ul>
                <button
    type="button"
    class="add-project-btn"
    onclick="addSuggestedProjectToResume(this)"
>
    我已完成该项目，加入项目经历
</button>
            </div>
            """

    awards_html = list_items(data.get("awards", []))
    others_html = list_items(data.get("others", []))
    skills_html = skills_tags(data.get("skills", []))

    return education_html, experience_html, projects_html, suggested_html, awards_html, others_html, skills_html


def render_resume_template_page(data, page_title="优化版简历"):
    education_html, experience_html, projects_html, suggested_html, awards_html, others_html, skills_html = build_resume_html(data)
    safe_plain_text = e(plain_text_resume(data))

    return f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>{e(page_title)}</title>
<style>
* {{ box-sizing: border-box; }}
body {{
    margin: 0;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, "Microsoft YaHei", sans-serif;
    background: #eef2f7;
    color: #111827;
}}
.page {{ padding: 40px 24px; }}
.container {{ max-width: 1120px; margin: auto; }}
.top {{
    display: flex; justify-content: space-between; align-items: center; margin-bottom: 24px;
}}
.back {{
    text-decoration: none; background: white; border: 1px solid #e5e7eb;
    color: #374151; padding: 10px 16px; border-radius: 12px; font-weight: 600;
}}
.action-card {{
    background: white; border-radius: 22px; padding: 24px;
    box-shadow: 0 12px 30px rgba(15, 23, 42, 0.08); margin-bottom: 24px;
}}
.action-row {{
    display: grid; grid-template-columns: 1fr 1fr 1fr 1fr; gap: 14px;
}}
.action-btn {{
    width: 100%; border: none; border-radius: 14px; padding: 14px;
    color: white; font-weight: 700; cursor: pointer;
}}
.pdf-btn {{ background: linear-gradient(135deg,#111827,#374151); }}
.word-btn {{ background: linear-gradient(135deg,#2563eb,#4f46e5); }}
.template-btn {{ background: linear-gradient(135deg,#16a34a,#22c55e); }}

.resume-wrapper {{
    background: white; border-radius: 24px; padding: 40px;
    box-shadow: 0 20px 45px rgba(15, 23, 42, 0.08);
}}
.resume-template {{ max-width: 850px; margin: 0 auto; background: white; }}
.resume-template[contenteditable="true"]:focus {{ outline: 3px solid #bfdbfe; }}

.resume-classic {{
    padding: 48px 56px; border: 1px solid #e5e7eb;
}}
.resume-classic .name {{
    font-size: 34px; font-weight: 800; text-align: center; margin-bottom: 8px;
}}
.resume-classic .contact {{
    text-align: center; color: #4b5563; margin-bottom: 28px; line-height: 1.8;
}}
.resume-classic .section-title {{
    font-size: 17px; font-weight: 800; border-bottom: 2px solid #111827;
    padding-bottom: 6px; margin-top: 26px; margin-bottom: 14px;
}}

.resume-notion {{
    display: grid; grid-template-columns: 260px 1fr; gap: 28px;
    padding: 32px; background: #fbfbfa; border: 1px solid #e5e7eb; border-radius: 22px;
}}
.sidebar {{
    background: #f1f5f9; border-radius: 20px; padding: 24px;
}}
.main {{
    background: white; border-radius: 20px; padding: 28px; border: 1px solid #e5e7eb;
}}
.resume-notion .name {{
    font-size: 30px; font-weight: 800; margin-bottom: 8px;
}}
.resume-notion .contact {{
    color: #64748b; font-size: 14px; line-height: 1.8; margin-bottom: 20px;
}}
.resume-notion .section-title {{
    font-size: 16px; font-weight: 800; margin-top: 24px; margin-bottom: 12px; color: #0f172a;
}}

.item {{ margin-bottom: 18px; }}
.item-title {{ font-weight: 800; font-size: 15px; margin-bottom: 4px; }}
.item-meta {{ color: #6b7280; font-size: 13px; margin-bottom: 8px; }}
ul {{ margin: 8px 0 0 20px; padding: 0; }}
li {{ margin-bottom: 7px; line-height: 1.65; }}
.skill-tag {{
    display: inline-block; background: #eef2ff; color: #3730a3;
    padding: 7px 11px; border-radius: 999px; margin: 4px; font-size: 13px; font-weight: 700;
}}
.skill-category {{ font-weight: 800; margin-top: 10px; color: #111827; }}
.summary {{ line-height: 1.8; color: #374151; }}
.hidden {{ display: none; }}
.muted {{ color: #9ca3af; }}
.warning-title {{ color: #b45309; border-color: #f59e0b !important; }}
.suggested {{
    background: #fffbeb; border: 1px solid #fde68a; border-radius: 12px; padding: 12px;
}}
.reason {{ color: #92400e; font-size: 13px; line-height: 1.6; }}
.project-subtitle {{
    font-weight: 800;
    color: #92400e;
    margin-top: 10px;
    margin-bottom: 4px;
    font-size: 13px;
}}
.add-project-btn {{
    margin-top: 10px;
    background: linear-gradient(135deg,#2563eb,#4f46e5);
    color: white;
    border: none;
    padding: 10px 14px;
    border-radius: 10px;
    cursor: pointer;
    font-size: 13px;
    font-weight: 700;
}}

.add-project-btn:hover {{
    opacity: 0.92;
}}
.added-project {{
    background: #f0fdf4;
    border: 1px solid #bbf7d0;
    border-radius: 10px;
    padding: 10px 12px;
}}
@media print {{
    body {{ background: white; }}
    .no-print, .top {{ display: none !important; }}
    .page {{ padding: 0; }}
    .container {{ max-width: 100%; }}
    .resume-wrapper {{ box-shadow: none; border-radius: 0; padding: 0; }}
    .resume-template {{ max-width: 100%; }}
}}
@media (max-width: 900px) {{
    .action-row {{ grid-template-columns: 1fr; }}
    .resume-notion {{ grid-template-columns: 1fr; }}
}}
</style>
</head>

<body>
<div class="page">
    <div class="container">

        <div class="top no-print">
            <h1>{e(page_title)}</h1>
            <a class="back" href="/">返回首页</a>
        </div>

        <div class="action-card no-print">
            <div class="action-row">
                <button class="action-btn template-btn" onclick="showTemplate('classic')">经典专业模板</button>
                <button class="action-btn template-btn" onclick="showTemplate('notion')">Notion 高级模板</button>
                <button class="action-btn pdf-btn" onclick="window.print()">导出 PDF</button>

                <form action="/export-word" method="post" onsubmit="syncWordContent()">
                    <input type="hidden" name="optimized_resume" id="wordContent" value="{safe_plain_text}">
                    <button type="submit" class="action-btn word-btn">导出 Word</button>
                </form>
            </div>
            <p style="color:#6b7280;font-size:13px;margin-bottom:0;">提示：简历区域可以直接点击修改；导出 Word 时会导出当前可见模板中的文本内容。</p>
        </div>

        <div class="resume-wrapper">
            <div id="classic-template" class="resume-template resume-classic" contenteditable="true">
                <div class="name">{e(data.get("name", "请补充姓名"))}</div>
                <div class="contact">{e(data.get("phone"))} ｜ {e(data.get("email"))} ｜ {e(data.get("city"))} ｜ {e(data.get("target_job"))}</div>

                <div class="section-title">个人简介</div>
                <div class="summary">{e(data.get("summary"))}</div>

                <div class="section-title">教育经历</div>
                {education_html}

                <div class="section-title">技能优势</div>
                <div>{skills_html}</div>

                <div class="section-title">实习经历 / 工作经历</div>
                {experience_html}

                <div class="section-title">项目经历</div>
                <div class="projects-container">
                    {projects_html}
                </div>

                {suggested_html}

                <div class="section-title">荣誉奖项</div>
                <ul>{awards_html}</ul>

                <div class="section-title">其他经历</div>
                <ul>{others_html}</ul>
            </div>

            <div id="notion-template" class="resume-template resume-notion hidden" contenteditable="true">
                <div class="sidebar">
                    <div class="name">{e(data.get("name", "请补充姓名"))}</div>
                    <div class="contact">
                        {e(data.get("phone"))}<br>
                        {e(data.get("email"))}<br>
                        {e(data.get("city"))}<br>
                        {e(data.get("target_job"))}
                    </div>

                    <div class="section-title">技能优势</div>
                    <div>{skills_html}</div>

                    <div class="section-title">教育经历</div>
                    {education_html}

                    <div class="section-title">荣誉奖项</div>
                    <ul>{awards_html}</ul>
                </div>

                <div class="main">
                    <div class="section-title">个人简介</div>
                    <div class="summary">{e(data.get("summary"))}</div>

                    <div class="section-title">实习经历 / 工作经历</div>
                    {experience_html}

                    <div class="section-title">项目经历</div>
                    <div class="projects-container">
                        {projects_html}
                    </div>

                    {suggested_html}

                    <div class="section-title">其他经历</div>
                    <ul>{others_html}</ul>
                </div>
            </div>
        </div>
    </div>
</div>

<script>
let activeTemplate = "classic";

function showTemplate(type) {{
    activeTemplate = type;
    const classic = document.getElementById("classic-template");
    const notion = document.getElementById("notion-template");

    if (type === "classic") {{
        classic.classList.remove("hidden");
        notion.classList.add("hidden");
    }} else {{
        notion.classList.remove("hidden");
        classic.classList.add("hidden");
    }}
}}

function syncWordContent() {{
    const active = document.getElementById(activeTemplate + "-template");
    document.getElementById("wordContent").value = active.innerText;
}}
function addSuggestedProjectToResume(button) {{
    const suggestedCard = button.closest(".suggested");

    // 优先把项目加入到当前按钮所在的模板中，避免 classic / notion 切换后找错位置
    const currentResume = button.closest(".resume-template");
    const projectsContainer = currentResume ? currentResume.querySelector(".projects-container") : null;

    if (!suggestedCard || !projectsContainer) {{
        alert("没有找到项目经历区域，请检查 projects-container 是否添加成功");
        return;
    }}

    const title = suggestedCard.querySelector(".item-title")?.innerText || "补做项目";
    const bulletLists = suggestedCard.querySelectorAll("ul");
    const resumeBullets = bulletLists[bulletLists.length - 1]?.innerHTML || "";

    const newProject = document.createElement("div");
    newProject.className = "item added-project";
    newProject.innerHTML =
        '<div class="item-title">' + title + '</div>' +
        '<div class="item-meta">已完成项目｜请补充时间</div>' +
        '<ul>' + resumeBullets + '</ul>';

    projectsContainer.appendChild(newProject);

    // 让用户立刻看到已经添加的位置
    newProject.scrollIntoView({{ behavior: "smooth", block: "center" }});

    button.innerText = "已加入项目经历";
    button.style.background = "#16a34a";
    button.disabled = true;
}}
</script>
</body>
</html>
"""

@app.get("/admin-codes", response_class=HTMLResponse)
def admin_codes(password: str = "", count: int = 10):
    if password != ADMIN_PASSWORD:
        return HTMLResponse("""
        <html>
        <head><meta charset="UTF-8"></head>
        <body style="font-family:Arial,'Microsoft YaHei';padding:40px;">
            <h2>无权限</h2>
            <p>后台密码错误。</p>
        </body>
        </html>
        """)

    new_codes = create_codes(count)

    codes_html = "".join(
        f"""
        <div style="padding:12px 16px;background:#f3f4f6;border-radius:10px;margin:8px 0;font-family:monospace;font-size:16px;">
            {code}
        </div>
        """
        for code in new_codes
    )

    return HTMLResponse(f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <title>生成激活码</title>
    </head>
    <body style="font-family:Arial,'Microsoft YaHei';background:#f3f4f6;padding:60px;">
        <div style="max-width:680px;margin:auto;background:white;padding:32px;border-radius:24px;box-shadow:0 20px 40px rgba(0,0,0,.08);">
            <h1>已生成 {count} 个一次性激活码</h1>
            <p>复制下面卡密发给付款用户。每个激活码只能使用一次。</p>

            {codes_html}

            <a href="/admin-codes?password={password}&count={count}"
               style="display:inline-block;margin-top:20px;background:#7c3aed;color:white;padding:12px 18px;border-radius:12px;text-decoration:none;font-weight:700;">
               再生成一批
            </a>
        </div>
    </body>
    </html>
    """)
@app.get("/", response_class=HTMLResponse)
def home():
    return RedirectResponse(url="/dashboard", status_code=302)


NAV_ITEMS = [
    {"path": "/dashboard", "label": "首页", "icon": "⌂"},
    {"path": "/resumes", "label": "我的简历", "icon": "▤"},
    {"path": "/analysis-history", "label": "简历分析记录", "icon": "◷"},
    {"path": "/project-library", "label": "项目库", "icon": "◇"},
    {"path": "/templates", "label": "模板库", "icon": "▦"},
    {"path": "/pro", "label": "PRO 会员", "icon": "★"},
]

MOCK_RESUMES = [
    {"title": "数据分析实习生简历", "role": "Data Analyst Intern", "score": 86, "updated": "今天 14:20", "status": "已优化"},
    {"title": "Java 后端开发简历", "role": "Java Developer", "score": 78, "updated": "昨天 21:10", "status": "待补充"},
    {"title": "产品经理校招简历", "role": "Product Manager", "score": 91, "updated": "05-22 09:34", "status": "可投递"},
]

MOCK_ANALYSES = [
    {"job": "字节跳动 数据分析实习生", "score": 88, "missing": "SQL 窗口函数、A/B Test", "time": "今天"},
    {"job": "腾讯 后端开发工程师", "score": 76, "missing": "Redis、高并发", "time": "昨天"},
    {"job": "小红书 用户运营实习生", "score": 82, "missing": "内容增长、用户分层", "time": "05-20"},
]

MOCK_PROJECTS = [
    {"name": "电商用户行为分析", "tags": ["Python", "SQL", "可视化"], "level": "适合数据岗", "desc": "从清洗、漏斗分析到转化建议，补齐业务分析项目表达。"},
    {"name": "校园二手交易小程序", "tags": ["Spring Boot", "MySQL", "Redis"], "level": "适合后端岗", "desc": "覆盖登录、商品发布、搜索、订单等可写入简历的核心模块。"},
    {"name": "AI 简历优化工具", "tags": ["FastAPI", "LLM", "Prompt"], "level": "适合 AI 应用岗", "desc": "展示大模型接口调用、文件解析和结构化生成能力。"},
]

MOCK_TEMPLATES = [
    {"name": "ATS 极简单栏", "type": "通用", "accent": "#4f46e5"},
    {"name": "技术岗项目强化", "type": "研发", "accent": "#0ea5e9"},
    {"name": "校招经历突出", "type": "学生", "accent": "#8b5cf6"},
    {"name": "产品运营双栏", "type": "运营", "accent": "#14b8a6"},
]

APP_CSS = """
:root {
    --primary: #4f46e5;
    --secondary: #7c3aed;
    --success: #22c55e;
    --ink: #111827;
    --muted: #64748b;
    --line: rgba(226,232,240,.9);
    --shadow: 0 24px 70px rgba(15,23,42,.08);
}
* { box-sizing: border-box; }
html { scroll-behavior: smooth; }
body {
    margin: 0;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, "Microsoft YaHei", sans-serif;
    background:
        radial-gradient(circle at 22% 8%, rgba(124,58,237,.13), transparent 30%),
        radial-gradient(circle at 86% 18%, rgba(79,70,229,.16), transparent 28%),
        linear-gradient(135deg,#f5f7ff 0%,#eef2ff 56%,#f8fbff 100%);
    color: var(--ink);
}
a { color: inherit; text-decoration: none; }
button, input, textarea { font-family: inherit; }
.app-shell { min-height: 100vh; display: grid; grid-template-columns: 220px minmax(0, 1fr) 320px; }
.sidebar { position: sticky; top: 0; height: 100vh; background: rgba(255,255,255,.72); backdrop-filter: blur(18px); border-right: 1px solid var(--line); padding: 22px 14px; display: flex; flex-direction: column; gap: 20px; z-index: 20; }
.logo { display: flex; align-items: center; gap: 10px; padding: 4px 6px; }
.logo-icon { width: 40px; height: 40px; border-radius: 14px; background: linear-gradient(135deg,var(--primary),var(--secondary)); color: white; display: flex; align-items: center; justify-content: center; font-weight: 900; box-shadow: 0 16px 32px rgba(79,70,229,.28); }
.logo-title { font-size: 18px; font-weight: 900; letter-spacing: .2px; }
.logo-sub { color: var(--muted); font-size: 12px; margin-top: 3px; }
.mobile-menu { display: none; position: fixed; top: 14px; left: 14px; z-index: 40; width: 42px; height: 42px; border-radius: 15px; color: white; background: linear-gradient(135deg,var(--primary),var(--secondary)); box-shadow: 0 18px 38px rgba(79,70,229,.25); }
.new-btn, button { border: 0; border-radius: 18px; cursor: pointer; font-weight: 900; transition: transform .18s ease, box-shadow .18s ease, opacity .18s ease; }
.new-btn { display: block; text-align: center; color: white; padding: 13px 14px; background: linear-gradient(135deg,var(--primary),var(--secondary)); box-shadow: 0 18px 38px rgba(79,70,229,.24); }
button:hover, .new-btn:hover, .primary-link:hover { transform: translateY(-2px); }
button:disabled { opacity: .72; cursor: wait; transform: none; }
.nav { display: grid; gap: 6px; }
.nav-item { display: flex; align-items: center; gap: 9px; padding: 11px 12px; border-radius: 16px; color: #526079; font-size: 14px; font-weight: 800; transition: .18s ease; }
.nav-item:hover { background: rgba(238,242,255,.8); transform: translateX(2px); }
.nav-item span:first-child { width: 20px; text-align: center; color: var(--secondary); }
.nav-item.active { background: #eef2ff; color: var(--primary); box-shadow: inset 0 0 0 1px rgba(79,70,229,.08); }
.side-card { margin-top: auto; background: linear-gradient(135deg,var(--primary),var(--secondary)); color: white; border-radius: 22px; padding: 16px; box-shadow: 0 18px 42px rgba(79,70,229,.28); }
.side-card h3 { margin: 0 0 8px; font-size: 16px; }
.side-card p { margin: 0 0 14px; font-size: 12px; line-height: 1.65; opacity: .92; }
.side-card a { display: block; text-align: center; background: white; color: var(--primary); padding: 10px; border-radius: 15px; font-weight: 900; }
.main { min-width: 0; padding: 28px 28px 48px; }
.main-inner { max-width: 1040px; margin: 0 auto; }
.topbar { display: none; }
h1 { margin: 0; font-size: 34px; line-height: 1.16; letter-spacing: 0; }
.hero-card { position: relative; overflow: hidden; border-radius: 28px; min-height: 360px; padding: 44px; margin-bottom: 18px; background: linear-gradient(135deg,rgba(255,255,255,.88),rgba(238,242,255,.72)); border: 1px solid rgba(255,255,255,.92); box-shadow: var(--shadow); isolation: isolate; animation: fadeUp .5s ease both; }
.hero-card::before { content: ""; position: absolute; inset: 0; background-image: linear-gradient(rgba(79,70,229,.07) 1px, transparent 1px), linear-gradient(90deg, rgba(79,70,229,.07) 1px, transparent 1px); background-size: 34px 34px; mask-image: radial-gradient(circle at 42% 30%, black, transparent 72%); opacity: .72; }
.hero-glow { position: absolute; width: 360px; height: 360px; right: -100px; top: -120px; background: radial-gradient(circle,rgba(124,58,237,.36),rgba(79,70,229,.14) 46%,transparent 72%); filter: blur(4px); animation: pulseGlow 5s ease-in-out infinite; }
.blob { position: absolute; border-radius: 999px; filter: blur(1px); opacity: .9; animation: floatSoft 7s ease-in-out infinite; }
.blob-one { width: 160px; height: 160px; right: 270px; bottom: -60px; background: radial-gradient(circle,rgba(34,197,94,.22),transparent 68%); }
.blob-two { width: 180px; height: 180px; left: 42%; top: -86px; background: radial-gradient(circle,rgba(79,70,229,.18),transparent 70%); animation-delay: -2s; }
.particle { position: absolute; width: 5px; height: 5px; border-radius: 99px; background: rgba(79,70,229,.45); box-shadow: 0 0 18px rgba(124,58,237,.45); animation: particleDrift 8s linear infinite; }
.p1 { left: 52%; top: 28%; } .p2 { left: 72%; top: 18%; animation-delay: -2s; } .p3 { left: 64%; top: 72%; animation-delay: -4s; }
.hero-content { position: relative; z-index: 2; max-width: 590px; }
.hero-kicker { display: inline-flex; align-items: center; gap: 8px; color: var(--primary); font-weight: 900; font-size: 13px; margin-bottom: 12px; padding: 8px 12px; border-radius: 999px; background: rgba(238,242,255,.82); border: 1px solid rgba(199,210,254,.8); }
.hero-kicker::before { content: ""; width: 8px; height: 8px; border-radius: 99px; background: var(--success); box-shadow: 0 0 0 5px rgba(34,197,94,.13); }
.hero-card h1 { font-size: 52px; line-height: 1.06; }
.hero-card p { color: #475569; font-size: 17px; line-height: 1.85; margin: 16px 0 0; }
.hero-actions { display: flex; flex-wrap: wrap; gap: 12px; margin-top: 24px; }
.hero-chip { color: #334155; background: rgba(255,255,255,.76); border: 1px solid var(--line); border-radius: 999px; padding: 10px 14px; font-weight: 800; font-size: 13px; }
.floating-resume { position: absolute; z-index: 2; right: 38px; top: 58px; width: 270px; background: rgba(255,255,255,.88); border: 1px solid rgba(255,255,255,.96); border-radius: 24px; padding: 18px; box-shadow: 0 28px 70px rgba(79,70,229,.18); backdrop-filter: blur(16px); animation: cardFloat 5s ease-in-out infinite; }
.scan-line { position: absolute; left: 16px; right: 16px; top: 74px; height: 2px; background: linear-gradient(90deg,transparent,var(--primary),var(--secondary),transparent); box-shadow: 0 0 22px rgba(124,58,237,.55); animation: scanMove 2.8s ease-in-out infinite; }
.float-row { height: 8px; border-radius: 99px; background: #e5e7eb; margin: 10px 0; overflow: hidden; }
.float-row span { display: block; height: 100%; border-radius: 99px; background: linear-gradient(90deg,var(--primary),var(--secondary)); }
.float-tags { display: flex; gap: 6px; flex-wrap: wrap; margin-top: 14px; }
.float-tags span { font-size: 11px; font-weight: 900; color: #4338ca; background: #eef2ff; padding: 5px 8px; border-radius: 999px; }
.trust-bar { display: flex; flex-wrap: wrap; gap: 10px; margin: 18px 0 22px; }
.trust-item { display: inline-flex; align-items: center; gap: 7px; padding: 9px 13px; border-radius: 999px; background: rgba(255,255,255,.74); border: 1px solid var(--line); color: #334155; font-size: 13px; font-weight: 800; box-shadow: 0 10px 26px rgba(15,23,42,.04); }
.trust-dot { width: 8px; height: 8px; border-radius: 999px; background: var(--success); box-shadow: 0 0 0 4px rgba(34,197,94,.12); }
.metric-row { display: none; }
.metric-card, .card, .panel, .table-card, .generate-card, .preview-card, .score-card, .diff-card { background: rgba(255,255,255,.78); border: 1px solid rgba(226,232,240,.88); border-radius: 24px; box-shadow: 0 20px 56px rgba(15,23,42,.07); backdrop-filter: blur(18px); transition: transform .2s ease, box-shadow .2s ease, border-color .2s ease; }
.card:hover, .panel:hover, .generate-card:hover, .preview-card:hover, .score-card:hover, .diff-card:hover { transform: translateY(-2px); box-shadow: 0 28px 70px rgba(15,23,42,.10); border-color: rgba(124,58,237,.18); }
.card, .panel, .table-card, .generate-card, .score-card, .diff-card { padding: 22px; }
.card h2, .panel h2, .table-card h2, .generate-card h2, .score-card h2, .diff-card h2 { margin: 0 0 8px; font-size: 20px; }
.card p, .panel p, .generate-card p, .score-card p, .diff-card p { color: var(--muted); line-height: 1.7; margin: 0 0 16px; font-size: 14px; }
.score-grid { display: grid; grid-template-columns: repeat(4, minmax(0, 1fr)); gap: 14px; margin: 18px 0; }
.score-card { position: relative; overflow: hidden; min-height: 150px; }
.score-card::before { content: ""; position: absolute; inset: 0; background: radial-gradient(circle at 80% 0%, rgba(124,58,237,.14), transparent 42%); pointer-events: none; }
.score-ring { --score: 80; width: 78px; height: 78px; border-radius: 50%; display: grid; place-items: center; margin-bottom: 12px; background: conic-gradient(var(--primary) calc(var(--score) * 1%), #e8edff 0); position: relative; }
.score-ring::after { content: ""; position: absolute; width: 58px; height: 58px; border-radius: 50%; background: rgba(255,255,255,.95); }
.score-ring strong { position: relative; z-index: 1; color: var(--primary); font-size: 18px; }
.score-label { position: relative; z-index: 1; font-size: 15px; font-weight: 900; }
.score-note { position: relative; z-index: 1; color: var(--muted); font-size: 12px; margin-top: 5px; }
.dashboard-grid { display: block; }
.form-shell { display: grid; gap: 16px; }
.form-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }
.step-card { position: relative; min-height: 100%; }
.step-head { display: flex; align-items: flex-start; gap: 13px; margin-bottom: 16px; }
.step-icon { width: 42px; height: 42px; flex: 0 0 42px; border-radius: 16px; display: flex; align-items: center; justify-content: center; color: white; font-weight: 900; background: linear-gradient(135deg,var(--primary),var(--secondary)); box-shadow: 0 14px 28px rgba(79,70,229,.20); }
.step-kicker { color: var(--primary); font-size: 12px; font-weight: 900; margin-bottom: 3px; }
.step-title { font-size: 19px; font-weight: 900; }
.step-desc { color: var(--muted); line-height: 1.65; font-size: 13px; margin-top: 4px; }
.upload-box { border: 2px dashed #c4b5fd; background: linear-gradient(135deg,#fbfaff,#eef2ff); border-radius: 24px; padding: 30px 18px; text-align: center; transition: .22s ease; position: relative; overflow: hidden; }
.upload-box::before { content: ""; position: absolute; inset: 0; background: linear-gradient(120deg,transparent,rgba(124,58,237,.08),transparent); transform: translateX(-100%); transition: .4s ease; }
.upload-box:hover, .upload-box.drag-over { border-color: var(--secondary); box-shadow: 0 0 0 6px rgba(124,58,237,.08), 0 24px 50px rgba(79,70,229,.10); transform: translateY(-1px); }
.upload-box:hover::before, .upload-box.drag-over::before { transform: translateX(100%); }
.upload-icon { font-size: 30px; margin-bottom: 10px; color: var(--primary); }
.file-input { position: absolute; opacity: 0; width: 1px; height: 1px; pointer-events: none; }
.upload-label { position: relative; z-index: 1; display: block; cursor: pointer; }
.file-meta { display: none; margin-top: 14px; padding: 10px 12px; border-radius: 16px; background: rgba(220,252,231,.86); color: #15803d; font-size: 13px; font-weight: 800; }
.file-meta.show { display: block; }
input[type="file"] { width: 100%; margin-top: 14px; color: #475569; }
textarea, input[type="text"] { width: 100%; border: 1px solid #cbd5e1; border-radius: 20px; padding: 15px; font-size: 15px; line-height: 1.65; outline: none; resize: vertical; background: rgba(255,255,255,.96); transition: .18s; }
textarea { min-height: 172px; }
textarea:focus, input[type="text"]:focus { border-color: var(--secondary); box-shadow: 0 0 0 5px rgba(124,58,237,.11); }
.extra-textarea { min-height: 118px; margin-top: 0; }
.hint { margin-top: 8px; font-size: 12px; color: #94a3b8; }
.suggestion-row { display: flex; flex-wrap: wrap; gap: 9px; margin-top: 14px; }
.suggestion-chip { border: 1px solid #ddd6fe; color: #5b21b6; background: rgba(245,243,255,.82); border-radius: 999px; padding: 8px 11px; font-size: 13px; font-weight: 800; cursor: pointer; transition: .18s ease; }
.suggestion-chip:hover { background: #ede9fe; transform: translateY(-1px); box-shadow: 0 10px 20px rgba(124,58,237,.10); }
.button-row { display: grid; grid-template-columns: 1fr 1fr; gap: 14px; margin-top: 16px; }
.btn-analysis, .btn-optimize, .btn-generate, .primary-link, .preview-btn { color: white; padding: 18px 20px; font-size: 16px; border-radius: 20px; box-shadow: 0 18px 38px rgba(79,70,229,.24); position: relative; overflow: hidden; }
.btn-analysis::after, .btn-optimize::after, .btn-generate::after, .primary-link::after { content: ""; position: absolute; inset: 0; background: linear-gradient(120deg,transparent,rgba(255,255,255,.28),transparent); transform: translateX(-120%); transition: .45s ease; }
.btn-analysis:hover::after, .btn-optimize:hover::after, .btn-generate:hover::after, .primary-link:hover::after { transform: translateX(120%); }
.btn-analysis, .primary-link { background: linear-gradient(135deg,#2563eb,var(--primary)); }
.btn-optimize, .btn-generate, .preview-btn { background: linear-gradient(135deg,var(--secondary),#a855f7); }
.btn-generate { width: 100%; margin-top: 14px; }
.action-note { text-align: center; color: var(--muted); font-size: 13px; margin-top: 12px; }
.generate-card { margin-top: 18px; position: relative; }
.pro-badge, .ats-badge { color: white; background: linear-gradient(135deg,var(--secondary),#a855f7); border-radius: 999px; padding: 6px 11px; font-size: 12px; font-weight: 900; }
.pro-badge { position: absolute; right: 20px; top: 20px; }
.loading { display: none; margin-top: 14px; text-align: center; color: var(--primary); font-weight: 900; }
.security { display: none; }
.diff-section { margin-top: 18px; }
.diff-head { display: flex; align-items: center; justify-content: space-between; gap: 12px; margin-bottom: 14px; }
.diff-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 16px; }
.diff-panel { border-radius: 22px; padding: 18px; background: rgba(248,250,252,.8); border: 1px solid var(--line); }
.diff-panel.after { background: linear-gradient(135deg,rgba(240,253,244,.88),rgba(238,242,255,.74)); border-color: rgba(34,197,94,.28); }
.diff-title { display: flex; justify-content: space-between; align-items: center; margin-bottom: 12px; font-weight: 900; }
.diff-score { color: white; border-radius: 999px; padding: 6px 10px; font-size: 12px; background: #94a3b8; }
.after .diff-score { background: var(--success); }
.diff-text { color: #475569; line-height: 1.85; font-size: 14px; }
.diff-added { display: inline; background: rgba(34,197,94,.16); color: #166534; border-radius: 7px; padding: 2px 5px; font-weight: 800; }
.lift-badge { color: #15803d; background: #dcfce7; border-radius: 999px; padding: 7px 11px; font-size: 12px; font-weight: 900; }
.right { padding: 28px 22px 28px 0; }
.preview-card { padding: 20px; margin-bottom: 18px; }
.preview-title { font-size: 18px; font-weight: 900; margin-bottom: 14px; }
.resume-paper { background: white; border: 1px solid #e5e7eb; border-radius: 22px; padding: 22px; box-shadow: 0 18px 40px rgba(15,23,42,.06); position: relative; overflow: hidden; }
.resume-paper::before { content: ""; position: absolute; top: 0; left: 0; right: 0; height: 4px; background: linear-gradient(90deg,var(--primary),var(--secondary),var(--success)); }
.resume-top { display: flex; align-items: flex-start; justify-content: space-between; gap: 10px; }
.resume-name { font-size: 24px; font-weight: 900; }
.resume-role { margin-top: 4px; color: var(--primary); font-weight: 800; }
.line { height: 1px; background: #e5e7eb; margin: 14px 0; }
.resume-paper p { color: #475569; font-size: 13px; line-height: 1.75; margin: 0 0 10px; }
.skill, .tag { display: inline-block; background: #eef2ff; color: #4338ca; padding: 6px 10px; border-radius: 999px; margin: 4px 4px 0 0; font-size: 12px; font-weight: 800; }
.success-tag { background: #dcfce7; color: #15803d; }
.preview-btn { width: 100%; margin-top: 14px; display: block; text-align: center; padding: 13px 16px; font-size: 14px; }
.list-grid { display: grid; grid-template-columns: repeat(3, minmax(0, 1fr)); gap: 16px; }
.item-title { font-size: 17px; font-weight: 900; margin-bottom: 6px; }
.item-meta { color: var(--muted); font-size: 13px; line-height: 1.65; }
.score-pill { display: inline-flex; align-items: center; justify-content: center; min-width: 58px; padding: 7px 10px; border-radius: 999px; background: #ecfeff; color: #0891b2; font-weight: 900; }
.table-card table { width: 100%; border-collapse: collapse; }
th, td { text-align: left; padding: 14px 10px; border-bottom: 1px solid #eef2f7; font-size: 14px; }
th { color: var(--muted); font-size: 12px; }
.template-swatch { height: 92px; border-radius: 18px; margin-bottom: 14px; background: linear-gradient(135deg, var(--accent), #eef2ff); }
.pro-hero { display: grid; grid-template-columns: 1.1fr .9fr; gap: 18px; align-items: stretch; }
.price { font-size: 42px; font-weight: 900; color: var(--primary); margin: 10px 0; }
.ai-loading-overlay { position: fixed; inset: 0; z-index: 100; display: none; align-items: center; justify-content: center; padding: 24px; background: rgba(15,23,42,.18); backdrop-filter: blur(16px); }
.ai-loading-overlay.active { display: flex; animation: fadeIn .22s ease both; }
.loading-modal { width: min(620px, 100%); border-radius: 28px; padding: 26px; background: linear-gradient(135deg,rgba(255,255,255,.86),rgba(238,242,255,.78)); border: 1px solid rgba(255,255,255,.88); box-shadow: 0 30px 100px rgba(79,70,229,.28); position: relative; overflow: hidden; }
.loading-modal::before { content: ""; position: absolute; width: 260px; height: 260px; right: -90px; top: -120px; background: radial-gradient(circle,rgba(124,58,237,.35),transparent 70%); animation: pulseGlow 4s ease-in-out infinite; }
.loading-orb { width: 58px; height: 58px; border-radius: 20px; display: grid; place-items: center; color: white; font-weight: 900; background: linear-gradient(135deg,var(--primary),var(--secondary)); box-shadow: 0 18px 44px rgba(79,70,229,.32); animation: cardFloat 2.6s ease-in-out infinite; }
.loading-title { position: relative; z-index: 1; margin: 18px 0 6px; font-size: 24px; font-weight: 900; }
.loading-subtitle { position: relative; z-index: 1; color: var(--primary); font-weight: 900; min-height: 24px; }
.progress-track { position: relative; z-index: 1; height: 10px; border-radius: 99px; background: rgba(226,232,240,.9); margin: 20px 0; overflow: hidden; }
.progress-fill { width: 42%; height: 100%; border-radius: 99px; background: linear-gradient(90deg,var(--primary),var(--secondary),var(--success)); animation: progressLoop 1.8s ease-in-out infinite; }
.skeleton-card { position: relative; z-index: 1; background: rgba(255,255,255,.72); border: 1px solid var(--line); border-radius: 22px; padding: 18px; }
.skeleton-line { height: 12px; border-radius: 999px; margin: 10px 0; background: linear-gradient(90deg,#e2e8f0,#f8fafc,#e2e8f0); background-size: 220% 100%; animation: shimmer 1.25s linear infinite; }
.skeleton-line.w70 { width: 70%; } .skeleton-line.w52 { width: 52%; } .skeleton-line.w90 { width: 90%; } .skeleton-line.w36 { width: 36%; }
@keyframes fadeIn { from { opacity: 0; } to { opacity: 1; } }
@keyframes fadeUp { from { opacity: 0; transform: translateY(12px); } to { opacity: 1; transform: translateY(0); } }
@keyframes pulseGlow { 0%,100% { transform: scale(1); opacity: .75; } 50% { transform: scale(1.08); opacity: 1; } }
@keyframes floatSoft { 0%,100% { transform: translate3d(0,0,0); } 50% { transform: translate3d(12px,-14px,0); } }
@keyframes particleDrift { 0% { transform: translateY(0); opacity: .2; } 35% { opacity: .8; } 100% { transform: translateY(-42px); opacity: .1; } }
@keyframes cardFloat { 0%,100% { transform: translateY(0); } 50% { transform: translateY(-8px); } }
@keyframes scanMove { 0%,100% { transform: translateY(0); opacity: .25; } 50% { transform: translateY(150px); opacity: 1; } }
@keyframes progressLoop { 0% { transform: translateX(-80%); width: 38%; } 50% { width: 62%; } 100% { transform: translateX(180%); width: 38%; } }
@keyframes shimmer { to { background-position: -220% 0; } }
@media (max-width: 1200px) { .app-shell { grid-template-columns: 220px minmax(0,1fr); } .right { display: block; grid-column: 2; padding: 0 28px 42px; } }
@media (max-width: 980px) { .floating-resume { position: relative; right: auto; top: auto; width: 100%; margin-top: 28px; } .hero-card { min-height: auto; } .score-grid { grid-template-columns: repeat(2, minmax(0, 1fr)); } }
@media (max-width: 860px) { .mobile-menu { display: block; } .app-shell { grid-template-columns: 1fr; } .sidebar { position: fixed; left: 12px; top: 64px; width: 220px; height: auto; max-height: calc(100vh - 82px); transform: translateX(-120%); transition: transform .22s ease; border-radius: 24px; box-shadow: 0 24px 70px rgba(15,23,42,.16); } body.sidebar-open .sidebar { transform: translateX(0); } .main { padding: 72px 14px 30px; } .right { grid-column: 1; padding: 0 14px 34px; } .hero-card { padding: 28px 22px; } .hero-card h1 { font-size: 36px; } .form-grid, .button-row, .list-grid, .pro-hero, .diff-grid { grid-template-columns: 1fr; } .score-grid { grid-template-columns: 1fr; } h1 { font-size: 28px; } }
"""


def render_nav(active_path):
    links = []
    for item in NAV_ITEMS:
        active = " active" if item["path"] == active_path else ""
        links.append(f'<a class="nav-item{active}" href="{item["path"]}"><span>{item["icon"]}</span><span>{item["label"]}</span></a>')
    return "".join(links)


def render_right_panel():
    return """
    <aside class="right">
        <div class="preview-card">
            <div class="preview-title">AI 简历效果预览</div>
            <div class="resume-paper">
                <div class="resume-top">
                    <div>
                        <div class="resume-name">张小明</div>
                        <div class="resume-role">Java 后端开发工程师</div>
                    </div>
                    <span class="ats-badge">ATS 86%</span>
                </div>
                <div class="line"></div>
                <p><strong>AI 优化标签</strong></p>
                <div>
                    <span class="skill success-tag">关键词强化</span>
                    <span class="skill success-tag">项目量化</span>
                    <span class="skill">表达精简</span>
                </div>
                <div class="line"></div>
                <p><strong>核心技能</strong></p>
                <div>
                    <span class="skill">Java</span>
                    <span class="skill">Spring Boot</span>
                    <span class="skill">MySQL</span>
                    <span class="skill">Redis</span>
                </div>
                <div class="line"></div>
                <p><strong>项目经历</strong></p>
                <p>校园二手交易小程序：负责商品发布、搜索筛选、订单流转和缓存优化，提升核心页面访问速度。</p>
                <a class="preview-btn" href="/dashboard#ai-compare">查看完整效果</a>
            </div>
        </div>
        <div class="preview-card">
            <div class="preview-title">实时工作流</div>
            <p style="color:#475569;line-height:1.8;margin:0;">上传简历后，AI 会解析 PDF、匹配 JD、补齐关键词，并生成可导出的优化版简历。</p>
        </div>
    </aside>
    """


def render_app_page(active_path, title, subtitle, content, eyebrow="AI Resume Dashboard"):
    page = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>__TITLE__｜AI Resume</title>
<style>__APP_CSS__</style>
</head>
<body>
<button class="mobile-menu" type="button" onclick="toggleSidebar()">☰</button>
<div class="app-shell">
    <aside class="sidebar">
        <a class="logo" href="/dashboard">
            <div class="logo-icon">AI</div>
            <div><div class="logo-title">AI Resume</div><div class="logo-sub">智能简历工作台</div></div>
        </a>
        <a class="new-btn" href="/dashboard">＋ 新建简历</a>
        <nav class="nav">__NAV__</nav>
        <div class="side-card">
            <h3>PRO 会员</h3>
            <p>解锁 AI 从零生成简历、项目推荐和更多模板能力。</p>
            <a href="/pro">立即升级</a>
        </div>
    </aside>
    <main class="main">
        <div class="topbar">
            <div><div class="eyebrow">__EYEBROW__</div><h1>__TITLE__</h1><p class="lead">__SUBTITLE__</p></div>
        </div>
        __CONTENT__
    </main>
    __RIGHT_PANEL__
</div>
<div class="ai-loading-overlay" id="aiLoadingOverlay" aria-live="polite" aria-hidden="true">
    <div class="loading-modal">
        <div class="loading-orb">AI</div>
        <div class="loading-title">AI 正在分析简历</div>
        <div class="loading-subtitle" id="loadingPhrase">正在解析 PDF...</div>
        <div class="progress-track"><div class="progress-fill"></div></div>
        <div class="skeleton-card">
            <div class="skeleton-line w36"></div>
            <div class="skeleton-line w70"></div>
            <div class="skeleton-line w90"></div>
            <div class="skeleton-line w52"></div>
        </div>
    </div>
</div>
<script>
let loadingTimer = null;
const loadingPhrases = [
    "正在解析 PDF...",
    "正在提取技能关键词...",
    "正在匹配岗位 JD...",
    "正在优化 ATS 通过率...",
    "正在补充项目经历...",
    "正在生成最终简历..."
];
function showLoading(event) {
    const submitter = event && event.submitter;
    if (submitter) {
        submitter.disabled = true;
        submitter.dataset.originalText = submitter.textContent;
        submitter.textContent = "AI 处理中...";
    }
    const overlay = document.getElementById("aiLoadingOverlay");
    const phrase = document.getElementById("loadingPhrase");
    if (overlay && phrase) {
        let index = 0;
        phrase.textContent = loadingPhrases[index];
        overlay.classList.add("active");
        overlay.setAttribute("aria-hidden", "false");
        clearInterval(loadingTimer);
        loadingTimer = setInterval(() => {
            index = (index + 1) % loadingPhrases.length;
            phrase.textContent = loadingPhrases[index];
        }, 1200);
    }
    return true;
}
function toggleSidebar() {
    document.body.classList.toggle("sidebar-open");
}
function formatFileSize(bytes) {
    if (!bytes) return "0 KB";
    if (bytes < 1024 * 1024) return Math.round(bytes / 1024) + " KB";
    return (bytes / 1024 / 1024).toFixed(1) + " MB";
}
function setupUploadBox() {
    const dropzone = document.getElementById("resumeDropzone");
    const input = document.getElementById("resumeFile");
    const meta = document.getElementById("resumeFileMeta");
    if (!dropzone || !input || !meta) return;
    const updateMeta = () => {
        const file = input.files && input.files[0];
        if (!file) return;
        meta.textContent = "✓ " + file.name + " · " + formatFileSize(file.size);
        meta.classList.add("show");
    };
    input.addEventListener("change", updateMeta);
    ["dragenter", "dragover"].forEach((name) => {
        dropzone.addEventListener(name, (e) => {
            e.preventDefault();
            dropzone.classList.add("drag-over");
        });
    });
    ["dragleave", "drop"].forEach((name) => {
        dropzone.addEventListener(name, (e) => {
            e.preventDefault();
            dropzone.classList.remove("drag-over");
        });
    });
    dropzone.addEventListener("drop", (e) => {
        const file = e.dataTransfer.files && e.dataTransfer.files[0];
        if (!file) return;
        const transfer = new DataTransfer();
        transfer.items.add(file);
        input.files = transfer.files;
        updateMeta();
    });
}
function setupSuggestionChips() {
    const textarea = document.getElementById("extraInfo");
    if (!textarea) return;
    document.querySelectorAll(".suggestion-chip").forEach((chip) => {
        chip.addEventListener("click", () => {
            const text = chip.dataset.suggestion || chip.textContent.trim();
            const prefix = textarea.value.trim() ? "\\n" : "";
            textarea.value += prefix + text;
            textarea.focus();
        });
    });
}
document.addEventListener("DOMContentLoaded", () => {
    setupUploadBox();
    setupSuggestionChips();
});
</script>
</body>
</html>
"""
    return (
        page
        .replace("__TITLE__", title)
        .replace("__SUBTITLE__", subtitle)
        .replace("__EYEBROW__", eyebrow)
        .replace("__APP_CSS__", APP_CSS)
        .replace("__NAV__", render_nav(active_path))
        .replace("__CONTENT__", content)
        .replace("__RIGHT_PANEL__", render_right_panel())
    )


def dashboard_content():
    return """
    <div class="main-inner">
        <section class="hero-card">
            <div class="hero-glow"></div>
            <div class="blob blob-one"></div>
            <div class="blob blob-two"></div>
            <span class="particle p1"></span>
            <span class="particle p2"></span>
            <span class="particle p3"></span>
            <div class="hero-content">
                <div class="hero-kicker">AI Resume Builder 正在待命</div>
                <h1>AI 简历生成器</h1>
                <p>上传简历 / 粘贴岗位 JD / 补充经历，AI 自动生成更适合投递的简历</p>
                <div class="hero-actions">
                    <span class="hero-chip">ATS 智能扫描</span>
                    <span class="hero-chip">JD 关键词理解</span>
                    <span class="hero-chip">项目经历增强</span>
                </div>
            </div>
            <div class="floating-resume" aria-label="floating resume demo">
                <div class="scan-line"></div>
                <strong>Resume Scan</strong>
                <div class="float-row"><span style="width:86%;"></span></div>
                <div class="float-row"><span style="width:64%;"></span></div>
                <div class="float-row"><span style="width:78%;"></span></div>
                <div class="float-tags"><span>Python</span><span>SQL</span><span>ATS 86%</span></div>
            </div>
        </section>

        <section class="trust-bar" aria-label="核心能力">
            <span class="trust-item"><span class="trust-dot"></span>ATS 优化</span>
            <span class="trust-item"><span class="trust-dot"></span>关键词匹配</span>
            <span class="trust-item"><span class="trust-dot"></span>PDF Word 导出</span>
            <span class="trust-item"><span class="trust-dot"></span>项目经历补全</span>
        </section>

        <section class="score-grid" aria-label="AI 简历评分">
            <div class="score-card"><div class="score-ring" style="--score:86;"><strong>86%</strong></div><div class="score-label">ATS 匹配度</div><div class="score-note">已覆盖核心 JD 关键词</div></div>
            <div class="score-card"><div class="score-ring" style="--score:78;"><strong>78%</strong></div><div class="score-label">项目竞争力</div><div class="score-note">建议补充量化结果</div></div>
            <div class="score-card"><div class="score-ring" style="--score:82;"><strong>82%</strong></div><div class="score-label">技能完整度</div><div class="score-note">技术栈表达较完整</div></div>
            <div class="score-card"><div class="score-ring" style="--score:74;"><strong>74%</strong></div><div class="score-label">面试机会指数</div><div class="score-note">优化后预计提升明显</div></div>
        </section>

        <section class="dashboard-grid">
            <div>
                <div class="card">
                    <form method="post" enctype="multipart/form-data" onsubmit="return showLoading(event)">
                        <div class="form-shell">
                            <div class="form-grid">
                                <div class="panel step-card">
                                    <div class="step-head">
                                        <div class="step-icon">1</div>
                                        <div>
                                            <div class="step-kicker">Step 1</div>
                                            <div class="step-title">上传简历</div>
                                            <div class="step-desc">拖拽或选择 PDF，AI 会自动提取教育、经历、项目和技能信息。</div>
                                        </div>
                                    </div>
                                    <div class="upload-box" id="resumeDropzone">
                                        <label class="upload-label" for="resumeFile">
                                            <div class="upload-icon">⇧</div>
                                            <strong>拖拽 PDF 到这里，或点击上传</strong>
                                            <div class="hint">推荐使用文字版 PDF，识别更稳定。</div>
                                            <input class="file-input" id="resumeFile" type="file" name="file" accept=".pdf" required>
                                            <div class="file-meta" id="resumeFileMeta">已选择文件</div>
                                        </label>
                                    </div>
                                </div>

                                <div class="panel step-card">
                                    <div class="step-head">
                                        <div class="step-icon">2</div>
                                        <div>
                                            <div class="step-kicker">Step 2</div>
                                            <div class="step-title">输入 JD</div>
                                            <div class="step-desc">粘贴目标岗位描述，AI 会识别关键词、能力要求和匹配差距。</div>
                                        </div>
                                    </div>
                                    <textarea name="jd" placeholder="请输入目标岗位描述，例如岗位职责、任职要求、技能要求等..."></textarea>
                                    <div style="margin-top:14px;">
                                        <strong>岗位截图上传（可选）</strong>
                                        <input type="file" name="jd_image" accept="image/*">
                                        <div class="hint">文字 JD 和截图 JD 会合并分析。</div>
                                    </div>
                                </div>
                            </div>

                            <div class="panel step-card">
                                <div class="step-head">
                                    <div class="step-icon">3</div>
                                    <div>
                                        <div class="step-kicker">Step 3</div>
                                        <div class="step-title">补充你想加入简历的内容（可选）</div>
                                        <div class="step-desc">把新增经历、亮点或作品链接放在这里，AI 会优先融入相关模块。</div>
                                    </div>
                                </div>
                                <textarea id="extraInfo" class="extra-textarea" name="extra_info" placeholder="例如：新项目、技能、证书、获奖、实习细节、作品链接等……"></textarea>
                                <div class="suggestion-row" aria-label="AI Suggestions">
                                    <button type="button" class="suggestion-chip" data-suggestion="数据分析项目：使用 Python / SQL 完成数据清洗、指标分析和可视化看板。">数据分析项目</button>
                                    <button type="button" class="suggestion-chip" data-suggestion="校园项目：负责需求拆解、核心功能实现和上线测试，沉淀完整项目文档。">校园项目</button>
                                    <button type="button" class="suggestion-chip" data-suggestion="实习经历：参与业务数据整理、用户反馈分析和跨部门协作，输出优化建议。">实习经历</button>
                                    <button type="button" class="suggestion-chip" data-suggestion="技能关键词：Python、SQL、Excel、Pandas、数据可视化、A/B Test。">技能关键词</button>
                                    <button type="button" class="suggestion-chip" data-suggestion="比赛获奖：参与校级/省级竞赛，负责方案设计、数据分析或系统实现。">比赛获奖</button>
                                    <button type="button" class="suggestion-chip" data-suggestion="开源项目：维护个人 GitHub 项目，包含 README、功能演示和持续迭代记录。">开源项目</button>
                                </div>
                            </div>
                        </div>

                        <div class="button-row">
                            <button class="btn-analysis" type="submit" formaction="/upload">开始 ATS 分析</button>
                            <button class="btn-optimize" type="submit" formaction="/optimize">一键生成优化版简历</button>
                        </div>
                        <div class="action-note">数据仅用于本次生成，不会保存或泄露。</div>
                        <div id="loading" class="loading">AI 正在分析中，请稍等...</div>
                    </form>
                </div>

                <section class="diff-section diff-card" id="ai-compare">
                    <div class="diff-head">
                        <div>
                            <h2>AI 优化效果对比</h2>
                            <p>从普通描述升级为更适合 ATS 和招聘方阅读的成果表达。</p>
                        </div>
                        <span class="lift-badge">ATS +43%</span>
                    </div>
                    <div class="diff-grid">
                        <div class="diff-panel">
                            <div class="diff-title">优化前 <span class="diff-score">43%</span></div>
                            <div class="diff-text">熟练使用 Python</div>
                        </div>
                        <div class="diff-panel after">
                            <div class="diff-title">优化后 <span class="diff-score">86%</span></div>
                            <div class="diff-text"><span class="diff-added">基于 Python 完成数据分析与自动化脚本开发</span>，熟悉 <span class="diff-added">Pandas / Numpy</span>，具备数据处理能力。</div>
                        </div>
                    </div>
                </section>

                <div class="generate-card" id="generate">
                    <div class="pro-badge">PRO</div>
                    <h2>AI 从零生成简历</h2>
                    <p>没有成型简历也可以使用。填写个人背景和目标岗位，AI 会生成完整简历并推荐可补做项目。</p>
                    <form action="/generate-resume" method="post" onsubmit="return showLoading(event)">
                        <textarea name="user_background" placeholder="请填写你的背景：姓名、学校、专业、学历、实习经历、项目经历、技能、获奖经历等..." required></textarea>
                        <textarea name="target_job" placeholder="请填写目标岗位，例如：数据分析实习生、产品经理、算法工程师、运营实习生..." style="min-height:96px;margin-top:12px;" required></textarea>
                        <input type="text" name="pro_code" autocomplete="off" placeholder="输入付费激活码（AI 从零生成需解锁）" style="margin-top:12px;">
                        <button class="btn-generate" type="submit">AI 从零生成简历</button>
                    </form>
                </div>
            </div>
        </section>
    </div>
    """

@app.get("/dashboard", response_class=HTMLResponse)
def dashboard():
    return render_app_page("/dashboard", "AI Resume Dashboard", "上传简历、分析 JD、生成优化版简历，所有核心工作流都在这里完成。", dashboard_content())


@app.get("/resumes", response_class=HTMLResponse)
def resumes_page():
    cards = "".join([f'<div class="card"><div class="item-title">{r["title"]}</div><div class="item-meta">目标岗位：{r["role"]}<br>更新时间：{r["updated"]}</div><p><span class="score-pill">{r["score"]}</span> <span class="tag">{r["status"]}</span></p></div>' for r in MOCK_RESUMES])
    return render_app_page("/resumes", "我的简历", "这里展示用户创建和优化过的简历。当前使用 mock 数据，后续可替换为数据库查询结果。", f'<section class="list-grid">{cards}</section>')


@app.get("/analysis-history", response_class=HTMLResponse)
def analysis_history_page():
    rows = "".join([f'<tr><td>{a["job"]}</td><td><span class="score-pill">{a["score"]}</span></td><td>{a["missing"]}</td><td>{a["time"]}</td></tr>' for a in MOCK_ANALYSES])
    content = f'<section class="table-card"><h2>分析记录</h2><table><thead><tr><th>目标岗位</th><th>ATS 分</th><th>待补关键词</th><th>时间</th></tr></thead><tbody>{rows}</tbody></table></section>'
    return render_app_page("/analysis-history", "简历分析记录", "沉淀每一次 ATS 分析结果，便于对比不同岗位和持续迭代简历。", content)


@app.get("/project-library", response_class=HTMLResponse)
def project_library_page():
    cards = []
    for p in MOCK_PROJECTS:
        tags = "".join([f'<span class="tag">{tag}</span>' for tag in p["tags"]])
        cards.append(f'<div class="card"><div class="item-title">{p["name"]}</div><div class="item-meta">{p["level"]}</div><p>{p["desc"]}</p><div>{tags}</div></div>')
    return render_app_page("/project-library", "项目库", "为不同岗位准备可落地、可写进简历的项目方向，后续可接入项目数据表。", f'<section class="list-grid">{"".join(cards)}</section>')


@app.get("/templates", response_class=HTMLResponse)
def templates_page():
    cards = "".join([f'<div class="card"><div class="template-swatch" style="--accent:{t["accent"]};"></div><div class="item-title">{t["name"]}</div><div class="item-meta">类型：{t["type"]}</div></div>' for t in MOCK_TEMPLATES])
    return render_app_page("/templates", "模板库", "选择适合岗位和经历阶段的简历模板。当前为静态展示，后续可绑定模板配置。", f'<section class="list-grid">{cards}</section>')


@app.get("/pro", response_class=HTMLResponse)
def pro_page():
    content = """
    <section class="pro-hero">
        <div class="card">
            <h2>解锁完整 AI Resume 能力</h2>
            <p>适合没有成型简历、需要从零生成、想要更多项目建议和模板能力的用户。</p>
            <div class="price">¥9.9 <span style="font-size:15px;color:#64748b;">/ 次体验</span></div>
            <a class="primary-link" style="display:inline-block;border-radius:16px;" href="/dashboard#generate">去使用 PRO 生成功能</a>
        </div>
        <div class="card">
            <h2>会员权益</h2>
            <p>AI 从零生成完整简历</p>
            <p>按目标岗位推荐补做项目</p>
            <p>优先使用高匹配度模板</p>
            <p>持续保存分析历史与优化版本</p>
        </div>
    </section>
    """
    return render_app_page("/pro", "PRO 会员", "升级后可使用 AI 从零生成简历，并获得更完整的项目和模板支持。", content)


@app.post("/upload", response_class=HTMLResponse)
async def upload_resume(
    file: UploadFile = File(...),
    jd: str = Form(""),
    extra_info: str = Form(""),
    jd_image: Optional[UploadFile] = File(None)
):
    resume_text = extract_pdf_text(file)
    jd_image_text = extract_image_text(jd_image)
    final_jd = jd + "\n" + jd_image_text

    client = get_deepseek_client()

    prompt = f"""
你是一名专业 ATS 简历优化专家。

请根据用户简历和目标岗位 JD，进行 ATS 匹配分析。

【简历内容】
{resume_text}

【岗位 JD】
{final_jd}

【用户补充信息】
{extra_info}

补充信息使用规则：
- 用户补充信息可以作为分析依据
- 如果补充信息明确包含真实经历、项目、技能、证书、获奖或实习细节，可以纳入优化建议
- 不允许把模糊信息过度包装成虚假经历
- 如果信息不完整，请提醒用户补充
- 补充信息要优先用于匹配岗位关键词和完善简历表达

请严格返回 JSON。
不要输出任何解释。
不要输出 markdown。
不要输出 ```json。
只返回纯 JSON。

JSON 格式如下：

{{
  "score": 0,
  "matched_keywords": [],
  "missing_keywords": [],
  "advantages": [],
  "problems": [],
  "suggestions": [],
  "optimized_experience": [],
  "summary": ""
}}

字段要求：
1. score 为 ATS 匹配分，范围 0-100
2. matched_keywords 为简历中已匹配岗位 JD 的关键词
3. missing_keywords 为岗位 JD 中要求但简历缺失的关键词
4. advantages 为简历优势
5. problems 为简历主要问题
6. suggestions 为具体优化建议
7. optimized_experience 为可优化后的简历表达示例
8. summary 为整体建议

重要规则：
- 不允许编造经历
- 不允许虚构技能
- 不允许新增公司、学历、项目
- 不允许虚构数据
- 不允许添加简历中完全没有出现过的技术栈
- 只能基于已有简历内容优化表达
- 可以强化已有经历中的关键词
- 如果岗位明显不匹配，要直接指出
- 输出必须专业
- 输出必须适合 ATS 系统
- 所有内容必须用中文
- 输出必须是合法 JSON
"""

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}]
    )

    ai_result = strip_json_fence(response.choices[0].message.content)

    try:
        data = json.loads(ai_result)
    except Exception:
        data = {
            "score": 0,
            "matched_keywords": [],
            "missing_keywords": [],
            "advantages": [],
            "problems": ["AI 返回格式解析失败，请重新分析。"],
            "suggestions": [],
            "optimized_experience": [],
            "summary": ai_result
        }

    score = data.get("score", 0)

    matched_html = render_tags(data.get("matched_keywords", []), "#dcfce7", "#166534")
    missing_html = render_tags(data.get("missing_keywords", []), "#fee2e2", "#991b1b")
    advantages_html = render_list(data.get("advantages", []))
    problems_html = render_list(data.get("problems", []))
    suggestions_html = render_list(data.get("suggestions", []))
    optimized_html = render_list(data.get("optimized_experience", []))
    summary_html = e(data.get("summary", ""))

    safe_resume_text = e(resume_text)
    safe_final_jd = e(final_jd)
    safe_extra_info = e(extra_info)

    return f"""
<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>AI ATS 分析结果</title>
<style>
* {{ box-sizing: border-box; }}
body {{
    margin: 0;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Arial, "Microsoft YaHei", sans-serif;
    background: #f3f4f6;
    color: #111827;
}}
.page {{ padding: 40px 24px; }}
.container {{ max-width: 1100px; margin: auto; }}
.top {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 24px; }}
h1 {{ margin: 0; font-size: 34px; }}
.back {{
    text-decoration: none; background: white; border: 1px solid #e5e7eb;
    color: #374151; padding: 10px 16px; border-radius: 12px;
}}
.score-card {{
    background: linear-gradient(135deg, #4f46e5, #2563eb); color: white;
    border-radius: 24px; padding: 32px; margin-bottom: 24px;
    box-shadow: 0 20px 45px rgba(37, 99, 235, 0.25);
}}
.score {{ font-size: 64px; font-weight: 800; margin: 10px 0; }}
.grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 24px; }}
.card {{
    background: white; border-radius: 22px; padding: 26px;
    box-shadow: 0 12px 30px rgba(15, 23, 42, 0.08); margin-bottom: 24px;
}}
.card h2 {{ margin-top: 0; font-size: 22px; }}
li {{ margin-bottom: 10px; line-height: 1.6; }}
.summary {{
    background: #eef2ff; border: 1px solid #c7d2fe;
    padding: 22px; border-radius: 18px; line-height: 1.8;
}}
pre {{
    white-space: pre-wrap; background: #f9fafb; padding: 18px;
    border-radius: 14px; max-height: 360px; overflow: auto; font-size: 13px;
}}
.btn-optimize {{
    width: 100%; padding: 18px; border: none; border-radius: 16px;
    background: linear-gradient(135deg,#16a34a,#22c55e);
    color: white; font-size: 18px; font-weight: 700; cursor: pointer;
}}
@media (max-width: 900px) {{
    .grid {{ grid-template-columns: 1fr; }}
    .top {{ display: block; }}
    .back {{ display: inline-block; margin-top: 16px; }}
}}
</style>
</head>

<body>
<div class="page">
    <div class="container">
        <div class="top">
            <h1>AI ATS 分析结果</h1>
            <a class="back" href="/">重新分析</a>
        </div>

        <div class="score-card">
            <div>ATS 匹配评分</div>
            <div class="score">{score} / 100</div>
            <div>分数越高，说明简历与目标岗位越匹配。</div>
        </div>

        <div class="grid">
            <div class="card"><h2>匹配关键词</h2>{matched_html}</div>
            <div class="card"><h2>缺失关键词</h2>{missing_html}</div>
        </div>

        <div class="grid">
            <div class="card"><h2>简历优势</h2>{advantages_html}</div>
            <div class="card"><h2>主要问题</h2>{problems_html}</div>
        </div>

        <div class="card"><h2>优化建议</h2>{suggestions_html}</div>
        <div class="card"><h2>可优化表达示例</h2>{optimized_html}</div>
        <div class="card"><h2>整体建议</h2><div class="summary">{summary_html}</div></div>

        <div class="card">
            <h2>一键生成优化版简历</h2>
            <p style="color:#6b7280;">基于当前简历和目标岗位，生成一版更适合投递的 ATS 优化简历。</p>
            <form action="/optimize" method="post">
                <textarea name="resume_text" style="display:none;">{safe_resume_text}</textarea>
                <textarea name="final_jd" style="display:none;">{safe_final_jd}</textarea>
                <textarea name="extra_info" style="display:none;">{safe_extra_info}</textarea>
                <button class="btn-optimize" type="submit">生成优化版简历</button>
            </form>
        </div>

        <div class="grid">
            <div class="card"><h2>解析出的简历内容</h2><pre>{safe_resume_text}</pre></div>
            <div class="card"><h2>岗位 JD 内容</h2><pre>{safe_final_jd}</pre></div>
        </div>
    </div>
</div>
</body>
</html>
"""


@app.post("/optimize", response_class=HTMLResponse)
async def optimize_resume(
    file: Optional[UploadFile] = File(None),
    jd: str = Form(""),
    extra_info: str = Form(""),
    jd_image: Optional[UploadFile] = File(None),
    resume_text: str = Form(""),
    final_jd: str = Form("")
):
    if file is not None and file.filename:
        resume_text = extract_pdf_text(file)
        jd_image_text = extract_image_text(jd_image)
        final_jd = jd + "\n" + jd_image_text

    client = get_deepseek_client()

    prompt = f"""
你是一名专业 ATS 简历优化专家和简历排版专家。

请根据原始简历和目标岗位 JD，生成一份结构化、真实、简洁、适合 ATS 的优化版简历。

【原始简历】
{resume_text}

【目标岗位 JD】
{final_jd}

【用户补充信息】
{extra_info}

补充信息使用规则：
- 用户补充信息可以作为简历内容来源
- 如果补充信息明确包含真实经历、项目、技能、证书、获奖或实习细节，可以合理加入简历
- 不允许把模糊信息包装成虚假公司、虚假岗位、虚假学历或虚假奖项
- 如果补充信息不完整，请写成“请补充”
- 补充信息要优先放到最相关模块：技能、项目经历、实习经历、荣誉证书或其他经历
- 如果补充信息与目标岗位相关，应优先强化到个人简介、技能优势和项目经历中

严格要求：
- 不允许编造经历
- 不允许新增公司、学历、项目
- 不允许虚构数据
- 不允许添加原简历中完全没有出现过的技术栈
- 只能基于已有经历优化表达
- 可以调整顺序，让内容更贴合目标岗位
- 可以强化已有经历中的关键词
- 输出中文
- 不要输出解释
- 不要输出 Markdown
- 不要输出 ```json
- 只返回合法 JSON

请严格返回以下 JSON 格式：

{{
  "name": "",
  "phone": "",
  "email": "",
  "city": "",
  "target_job": "",
  "summary": "",
  "education": [
    {{
      "school": "",
      "major": "",
      "degree": "",
      "time": "",
      "details": []
    }}
  ],
  "skills": [],
  "experiences": [
    {{
      "company": "",
      "role": "",
      "time": "",
      "bullets": []
    }}
  ],
  "projects": [
    {{
      "name": "",
      "role": "",
      "time": "",
      "bullets": []
    }}
  ],
  "suggested_projects": [],
  "awards": [],
  "others": []
}}
"""

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}]
    )

    ai_result = strip_json_fence(response.choices[0].message.content)

    try:
        data = json.loads(ai_result)
    except Exception:
        data = {
            "name": "优化版简历",
            "phone": "",
            "email": "",
            "city": "",
            "target_job": "",
            "summary": ai_result,
            "education": [],
            "skills": [],
            "experiences": [],
            "projects": [],
            "suggested_projects": [],
            "awards": [],
            "others": []
        }

    return render_resume_template_page(data, "优化版简历")


@app.post("/generate-resume", response_class=HTMLResponse)
async def generate_resume(
    user_background: str = Form(...),
    target_job: str = Form(...),
    pro_code: str = Form("")
):
    if not is_pro_user(pro_code):
        return HTMLResponse("""
        <html>
        <head>
        <meta charset="UTF-8">
        <title>需要付费解锁</title>
        </head>
        <body style="font-family:Arial,'Microsoft YaHei';background:#f3f4f6;padding:60px;">
            <div style="max-width:560px;margin:auto;background:white;padding:36px;border-radius:24px;box-shadow:0 20px 40px rgba(0,0,0,.08);">
                <h1>AI 从零生成简历需付费解锁</h1>
                <p>该功能适合还没有成型简历的用户，会根据你的背景和目标岗位生成完整简历，并推荐可补做项目。</p>
                <h2 style="color:#7c3aed;">体验价：¥9.9 / 次</h2>
                <p>付款后获取激活码，即可使用。</p>
                <div style="background:#f9fafb;border:1px dashed #d1d5db;padding:24px;border-radius:16px;text-align:center;margin:20px 0;">
    <img src="/static/pay.jpg"
         style="max-width:240px;border-radius:12px;">
</div>
<p style="
    margin-top:16px;
    color:#6b7280;
    text-align:center;
    line-height:1.8;
    font-size:14px;
">
付款后请添加微信：DDC_321<br>
发送付款截图领取一次性激活码
</p>
                <a href="/" style="display:block;text-align:center;background:#7c3aed;color:white;text-decoration:none;padding:14px;border-radius:14px;font-weight:700;">返回首页输入激活码</a>
            </div>
        </body>
        </html>
        """)
    client = get_deepseek_client()

    prompt = f"""
你是一名专业 ATS 简历专家。

请根据用户背景和目标岗位，生成结构化简历数据。

【用户背景】
{user_background}

【目标岗位】
{target_job}

重要原则：
- 用户明确提供过的经历，可以直接优化
- 用户没有提供的公司、学校、学历、奖项、真实工作经历，不允许编造
- 如果目标岗位需要某些能力，但用户没有相关经历，可以生成“示例项目”或“可补充项目”
- 所有 AI 补充的内容必须标注为“示例项目，需用户确认后使用”
- 不允许把示例项目伪装成真实经历
- 不允许虚构公司名称、任职时间、学历、奖项
- 可以根据岗位需求生成学习项目、个人项目、作品集项目、课程项目
- 用户没有提供的信息，写成“请补充”
- 输出中文
- 只返回合法 JSON
- 不要 markdown
- 不要解释
- 控制在一页简历长度
- 每段经历最多 3 条 bullet
- 每条 bullet 不超过 35 个中文字符
- 不要生成长篇解释
- 只保留最重要经历
- 优先保留与岗位最相关内容
- 必须推荐 2-3 个适合目标岗位的补做项目
- 推荐项目必须是用户可以在 3-7 天内真实完成的项目
- 推荐项目不能伪装成真实经历
- 推荐项目只放在 suggested_projects 字段
- 不允许把推荐项目放进 projects 字段
- must_complete_tasks 写用户真实需要完成的任务
- resume_bullets_after_completion 写完成后才能放进简历的表达

JSON 格式如下：

{{
  "name": "",
  "target_job": "",
  "phone": "",
  "email": "",
  "city": "",
  "summary": "",
  "education": [
    {{
      "school": "",
      "major": "",
      "degree": "",
      "time": "",
      "details": []
    }}
  ],
  "skills": [
    {{
      "category": "",
      "items": []
    }}
  ],
  "experiences": [
    {{
      "company": "",
      "role": "",
      "time": "",
      "bullets": []
    }}
  ],
  "projects": [
    {{
      "name": "",
      "role": "",
      "time": "",
      "bullets": []
    }}
  ],
  "suggested_projects": [
    {{
  "name": "",
  "suitable_for": "",
  "difficulty": "",
  "estimated_days": "",
  "reason": "",
  "must_complete_tasks": [],
  "resume_bullets_after_completion": []
}}
],
  "awards": [],
  "others": []
}}
"""

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}]
    )

    ai_result = strip_json_fence(response.choices[0].message.content)

    try:
        data = json.loads(ai_result)
    except Exception:
        data = {
            "name": "请补充姓名",
            "target_job": target_job,
            "phone": "请补充手机号",
            "email": "请补充邮箱",
            "city": "请补充城市",
            "summary": ai_result,
            "education": [],
            "skills": [],
            "experiences": [],
            "projects": [],
            "suggested_projects": [],
            "awards": [],
            "others": []
        }

    return render_resume_template_page(data, "AI 从零生成简历")


@app.post("/export-word")
async def export_word(optimized_resume: str = Form(...)):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    def set_cell_bg(cell, color):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tc_pr.append(shd)

    def set_cell_border(cell, color="E5E7EB"):
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ["top", "left", "bottom", "right"]:
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), "single")
            tag.set(qn("w:sz"), "6")
            tag.set(qn("w:space"), "0")
            tag.set(qn("w:color"), color)
            borders.append(tag)
        tc_pr.append(borders)

    def add_run(paragraph, text, size=9.5, bold=False, color="111827"):
        run = paragraph.add_run(text)
        run.bold = bold
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        return run

    def add_section_title(cell, title):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, title, size=11, bold=True, color="2563EB")

    def add_bullet(cell, text):
        p = cell.add_paragraph(style=None)
        p.paragraph_format.left_indent = Pt(10)
        p.paragraph_format.first_line_indent = Pt(-10)
        p.paragraph_format.line_spacing = 1.12
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "• ", size=9, color="2563EB")
        add_run(p, text, size=9, color="374151")

    def clean_line(line):
        return line.strip().replace("**", "").replace("###", "").replace("---", "")

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(9.5)

    lines = [clean_line(x) for x in optimized_resume.split("\n") if clean_line(x)]

    name = lines[0] if lines else "优化版简历"
    rest_lines = lines[1:]

    # 顶部姓名
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, name, size=20, bold=True, color="111827")

    # 联系方式
    if rest_lines:
        contact = rest_lines[0]
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(10)
        add_run(p, contact, size=8.5, color="6B7280")
        rest_lines = rest_lines[1:]

    # 主体双栏表格
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    left = table.rows[0].cells[0]
    right = table.rows[0].cells[1]

    left.width = Inches(2.0)
    right.width = Inches(4.6)

    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    set_cell_bg(left, "F3F4F6")
    set_cell_border(left, "FFFFFF")
    set_cell_border(right, "FFFFFF")

    # 分配内容：技能/教育/奖项放左边，经历/项目/简介放右边
    current_cell = right
    current_title = ""

    left_titles = ["技能", "技能优势", "教育经历", "荣誉奖项", "证书", "其他"]
    right_titles = ["个人简介", "求职方向", "实习经历", "工作经历", "项目经历", "校园经历"]

    for line in rest_lines:
        is_title = any(t in line for t in left_titles + right_titles) and len(line) <= 20

        if is_title:
            current_title = line
            if any(t in line for t in left_titles):
                current_cell = left
            else:
                current_cell = right
            add_section_title(current_cell, line)
            continue

        if line.startswith("-") or line.startswith("•"):
            add_bullet(current_cell, line.lstrip("-• ").strip())
        else:
            p = current_cell.add_paragraph()
            p.paragraph_format.line_spacing = 1.12
            p.paragraph_format.space_after = Pt(2)

            # 经历/项目标题加粗
            if "｜" in line or "|" in line:
                add_run(p, line, size=9.2, bold=True, color="111827")
            else:
                add_run(p, line, size=9, color="374151")

    file_name = f"optimized_resume_{uuid.uuid4().hex}.docx"
    doc.save(file_name)

    return FileResponse(
        path=file_name,
        filename="优化版简历.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.post("/export-word-html")
async def export_word_html(resume_html: str = Form(...)):
    full_html = f"""
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: "Microsoft YaHei", Arial, sans-serif;
                color: #111827;
            }}
            .resume-paper {{
                width: 760px;
                margin: 0 auto;
                padding: 40px;
            }}
            .name {{
                font-size: 28px;
                font-weight: bold;
                text-align: center;
            }}
            .contact {{
                text-align: center;
                color: #666;
                margin-bottom: 20px;
            }}
            .section-title {{
                font-size: 16px;
                font-weight: bold;
                color: #2563eb;
                border-bottom: 1px solid #111827;
                padding-bottom: 4px;
                margin-top: 18px;
                margin-bottom: 10px;
            }}
            .skill-tag {{
                display: inline-block;
                background: #eef2ff;
                color: #3730a3;
                padding: 4px 8px;
                border-radius: 10px;
                margin: 3px;
                font-size: 12px;
            }}
            li {{
                margin-bottom: 5px;
                line-height: 1.5;
            }}
        </style>
    </head>
    <body>
        <div class="resume-paper">
            {resume_html}
        </div>
    </body>
    </html>
    """

    file_name = f"resume_{uuid.uuid4().hex}.doc"

    with open(file_name, "w", encoding="utf-8") as f:
        f.write(full_html)

    return FileResponse(
        path=file_name,
        filename="优化版简历.doc",
        media_type="application/msword"
    )



